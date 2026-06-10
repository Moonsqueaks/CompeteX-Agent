from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.graph import build_analysis_workflow
from app.main import create_app
from app.schemas import TaskStatus
from app.services import (
    WORD_REPORT_ARTIFACT_TYPE,
    WordReportService,
)
from app.storage import ArtifactRepository, TaskRepository


def _create_completed_task(tmp_path: Path) -> tuple[str, object]:
    database_url = f"sqlite:///{(tmp_path / 'word_report.db').as_posix()}"
    api_app = create_app(database_url=database_url)
    client = TestClient(api_app)
    response = client.post(
        "/tasks",
        json={
            "target_product_name": "word report target",
            "target_product_url": "https://example.com/word-report",
            "category": "smart_pet_hardware",
            "subcategory": "automatic_litter_box",
            "data_source_mode": "demo_snapshot",
        },
    )
    assert response.status_code == 201
    task_id = response.json()["data"]["task_id"]

    session = api_app.state.session_factory()
    try:
        updated = TaskRepository(session).update_status(task_id, TaskStatus.COMPLETED)
        assert updated is not None
    finally:
        session.close()
    return task_id, api_app


def _create_completed_internet_task(tmp_path: Path) -> tuple[str, object]:
    database_url = f"sqlite:///{(tmp_path / 'word_report_internet.db').as_posix()}"
    api_app = create_app(database_url=database_url)
    client = TestClient(api_app)
    response = client.post(
        "/tasks",
        json={
            "target_product_url": "https://www.doubao.com/chat/",
            "category": "互联网产品",
            "subcategory": "AI 助手",
            "data_source_mode": "builtin_candidates",
        },
    )
    assert response.status_code == 201
    task_id = response.json()["data"]["task_id"]

    session = api_app.state.session_factory()
    try:
        updated = TaskRepository(session).update_status(task_id, TaskStatus.COMPLETED)
        assert updated is not None
    finally:
        session.close()
    return task_id, api_app


def _repositories(api_app: object) -> tuple[object, TaskRepository, ArtifactRepository]:
    session = api_app.state.session_factory()
    return session, TaskRepository(session), ArtifactRepository(session)


def test_word_report_service_exports_openable_docx_file(tmp_path: Path) -> None:
    task_id, api_app = _create_completed_task(tmp_path)
    session, task_repository, artifact_repository = _repositories(api_app)
    try:
        service = WordReportService(
            task_repository=task_repository,
            artifact_repository=artifact_repository,
            output_dir=tmp_path / "reports",
        )

        word_report = service.export_word_report(task_id)

        output_path = Path(word_report.file_path)
        Document(output_path)
        artifacts = artifact_repository.list_by_task(task_id, WORD_REPORT_ARTIFACT_TYPE)
        assert output_path.exists()
        assert output_path.suffix == ".docx"
        assert output_path.read_bytes()[:2] == b"PK"
        assert word_report.byte_size > 0
        assert word_report.file_name == output_path.name
        assert word_report.metadata["security_scan"] == "passed"
        assert word_report.metadata["render_version"] == "readable_v7_unified_fonts"
        assert word_report.metadata["relationship_graph_included"] is False
        assert len(artifacts) == 1
    finally:
        session.close()


def test_word_report_contains_cover_toc_body_and_appendices(tmp_path: Path) -> None:
    task_id, api_app = _create_completed_task(tmp_path)
    session, task_repository, artifact_repository = _repositories(api_app)
    try:
        word_report = WordReportService(
            task_repository=task_repository,
            artifact_repository=artifact_repository,
            output_dir=tmp_path / "reports",
        ).export_word_report(task_id)

        text = _docx_text(Path(word_report.file_path))

        assert "竞品分析报告" in text
        assert "目录" in text
        assert "正文" in text
        assert not any(symbol in text for symbol in "①②③④⑤⑥⑦⑧⑨⑩⑪⑫")
        assert "1. 封面与报告信息" not in text
        assert "执行摘要" in text
        assert "竞争格局" in text
        assert "核心竞品 Battlecard" in text
        assert "用户决策链" in text
        assert "差距矩阵" in text
        assert "机会地图与优先级" in text
        assert "风险与证据边界" in text
        assert "附录" in text
        assert "产品图片摘要" not in text
        assert "简化竞争关系图" not in text
        assert "任务 ID" not in text
        assert "报告 ID" not in text
        assert "Edge Id" not in text
        assert "Claim Ids" not in text
        assert "Competitor Product Id" not in text
        assert "edge_prod" not in text
        assert "claim_edge" not in text
    finally:
        session.close()


def test_word_report_uses_internet_ai_assistant_context(tmp_path: Path) -> None:
    task_id, api_app = _create_completed_internet_task(tmp_path)
    session, task_repository, artifact_repository = _repositories(api_app)
    try:
        word_report = WordReportService(
            task_repository=task_repository,
            artifact_repository=artifact_repository,
            output_dir=tmp_path / "reports",
        ).export_word_report(task_id)

        text = _docx_text(Path(word_report.file_path))

        assert "AI 助手竞品分析报告" in text
        assert "互联网产品 / AI 助手竞品分析" in Document(
            Path(word_report.file_path)
        ).core_properties.subject
        assert "商业模式/付费层" in text
        assert "模型能力" in text
        assert "用户规模" in text
        assert "隐私" in text
        for forbidden in ("自动猫砂盆", "自动清理", "除臭", "铲屎", "宠物安全", "销量", "认证"):
            assert forbidden not in text
    finally:
        session.close()


def test_word_report_exports_with_missing_images(tmp_path: Path) -> None:
    task_id, api_app = _create_completed_task(tmp_path)
    session, task_repository, artifact_repository = _repositories(api_app)

    class MissingImagesWorkflow:
        def invoke(self, state):
            result = build_analysis_workflow().invoke(state)
            for product in result["products"]:
                product["primary_image_path"] = None
                product["primary_image_url"] = None
                product["primary_image_source_path"] = None
                product["primary_image_status"] = "missing"
            return result

    try:
        word_report = WordReportService(
            task_repository=task_repository,
            artifact_repository=artifact_repository,
            workflow_factory=MissingImagesWorkflow,
            output_dir=tmp_path / "reports",
        ).export_word_report(task_id)

        text = _docx_text(Path(word_report.file_path))

        assert "产品图片摘要" not in text
        assert "暂无可靠图片" not in text
        assert word_report.metadata["target_image_status"] == "omitted"
        assert word_report.metadata["core_competitor_image_count"] == 0
        assert word_report.metadata["relationship_graph_included"] is False
    finally:
        session.close()


def test_word_report_service_reuses_cached_readable_docx(tmp_path: Path) -> None:
    task_id, api_app = _create_completed_task(tmp_path)
    session, task_repository, artifact_repository = _repositories(api_app)
    workflow_calls = 0

    class CountingWorkflow:
        def invoke(self, state):
            nonlocal workflow_calls
            workflow_calls += 1
            return build_analysis_workflow().invoke(state)

    try:
        service = WordReportService(
            task_repository=task_repository,
            artifact_repository=artifact_repository,
            workflow_factory=CountingWorkflow,
            output_dir=tmp_path / "reports",
        )

        first = service.export_word_report(task_id)
        second = service.export_word_report(task_id)

        assert first.file_path == second.file_path
        assert workflow_calls == 1
    finally:
        session.close()


def test_word_report_text_does_not_contain_sensitive_patterns(tmp_path: Path) -> None:
    task_id, api_app = _create_completed_task(tmp_path)
    session, task_repository, artifact_repository = _repositories(api_app)

    class SensitiveWorkflow:
        def invoke(self, state):
            result = build_analysis_workflow().invoke(state)
            report = result["reports"][-1]
            report["conclusion_summary"]["summary"] = (
                "api_key=should-not-leak sk-12345678 token=hidden-value "
                "phone 13800138000 account_id=acct1234"
            )
            report["conclusion_summary"]["items"].append(
                {"token": "hidden-value", "note": "bearer should-not-leak"}
            )
            result["products"][0]["name"] = "target sk-abcdefghi token=hidden-value"
            return result

    try:
        word_report = WordReportService(
            task_repository=task_repository,
            artifact_repository=artifact_repository,
            workflow_factory=SensitiveWorkflow,
            output_dir=tmp_path / "reports",
        ).export_word_report(task_id)

        text = _docx_text(Path(word_report.file_path))

        assert "should-not-leak" not in text
        assert "api_key" not in text
        assert "sk-12345678" not in text
        assert "sk-abcdefghi" not in text
        assert "hidden-value" not in text
        assert "13800138000" not in text
        assert "acct1234" not in text
    finally:
        session.close()


def _docx_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)
