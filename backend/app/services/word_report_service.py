import re
from collections.abc import Callable, Mapping, Sequence
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from urllib.parse import unquote

from docx import Document
from docx.shared import Inches

from app.graph import build_analysis_workflow, create_initial_state
from app.schemas import (
    AnalysisTask,
    BattlefieldData,
    Product,
    ProductRole,
    RelationshipGraphImage,
    ReportData,
    ReportSection,
    TaskStatus,
    TraceData,
    WordReport,
)
from app.schemas.common import JsonObject
from app.security import contains_sensitive_text, redact_sensitive_text, redact_sensitive_value
from app.services.markdown_renderer import DEFAULT_REPORTS_DIR
from app.services.relationship_graph_service import (
    RELATIONSHIP_GRAPH_ARTIFACT_TYPE,
    render_relationship_graph_png,
)
from app.services.report_service import REPORT_ARTIFACT_TYPE, redact_report_data
from app.services.trace_service import TRACE_ARTIFACT_TYPE, _build_trace_data, _trace_artifact_id
from app.storage import ArtifactRepository, TaskRepository

WORD_REPORT_ARTIFACT_TYPE = "word_report"
WORD_REPORT_RENDER_VERSION = "readable_v4"
NO_RELIABLE_IMAGE = "暂无可靠图片"
NO_RELIABLE_DATA = "暂无可靠数据"
_WORD_REPORT_READABLE_STATUSES = {TaskStatus.COMPLETED, TaskStatus.HUMAN_REVIEWING}
_SAFE_FILE_STEM_PATTERN = re.compile(r"[^A-Za-z0-9_.-]+")
_PROJECT_ROOT = Path(__file__).resolve().parents[3]
_RAW_ASSETS_DIR = _PROJECT_ROOT / "data" / "raw"
_RAW_ASSET_URL_PREFIX = "/assets/raw/"
_SUPPORTED_DOCX_IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpg", ".jpeg", ".png"}
_MAX_TEXT_CHARS = 600
_MAX_RENDERED_ITEMS = 18
_MAX_WORD_SECTION_ITEMS = 5
_INTERNAL_REPORT_KEYS = {
    "appendix_type",
    "basis_edge_id",
    "battlecard_id",
    "claim_id",
    "claim_ids",
    "claims",
    "collection_repair",
    "competitor_product_id",
    "edge_id",
    "edge_ids",
    "gap_id",
    "evidence_id",
    "evidence_ids",
    "items",
    "opportunity_id",
    "product_id",
    "qa_agent",
    "report_id",
    "review_insight_id",
    "screenshot_path",
    "source_url",
    "task_id",
}
_SECTION_TITLE_OVERRIDES = {
    "conclusion_summary": "结论摘要",
    "competitive_landscape_judgment": "竞争格局判断",
    "core_competitor_analysis": "核心竞品拆解",
    "user_decision_chain_analysis": "用户决策链分析",
    "target_opportunities_and_risks": "目标产品机会与风险",
    "product_strategy_recommendations": "产品策略建议",
    "evidence_quality_appendix": "证据与质检附录",
    "analysis_process_appendix": "分析流程与系统能力附录",
}

WorkflowFactory = Callable[[], Any]
GraphRenderer = Callable[..., Path]


class WordRenderError(ValueError):
    pass


class WordReportServiceError(Exception):
    def __init__(
        self,
        code: str,
        message: str,
        *,
        status_code: int,
        details: dict[str, Any] | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.message = message
        self.status_code = status_code
        self.details = details or {}


class WordReportService:
    def __init__(
        self,
        *,
        task_repository: TaskRepository,
        artifact_repository: ArtifactRepository,
        workflow_factory: WorkflowFactory = build_analysis_workflow,
        output_dir: Path | str | None = None,
        graph_renderer: GraphRenderer | None = None,
    ) -> None:
        self.task_repository = task_repository
        self.artifact_repository = artifact_repository
        self.workflow_factory = workflow_factory
        self.output_dir = output_dir
        self.graph_renderer = graph_renderer or render_relationship_graph_png

    def export_word_report(self, task_id: str) -> WordReport:
        task = self._get_completed_task(task_id)

        report = self._latest_cached_report(task.task_id)
        if report is None:
            state = self._run_workflow(task)
            report = self._report_from_state(task, state)

        cached_word_report = self._cached_word_report(task.task_id, report.report_id)
        if cached_word_report is not None:
            return cached_word_report

        try:
            word_report = render_word_report(
                report,
                output_dir=self.output_dir,
                extra_metadata={"render_version": WORD_REPORT_RENDER_VERSION},
            )
        except Exception as exc:
            self._record_word_export_failure(
                task_id=task.task_id,
                report=report,
                phase="docx_render_or_write",
                exc=exc,
            )
            raise WordReportServiceError(
                "WORD_REPORT_EXPORT_FAILED",
                "Word report export failed",
                status_code=500,
                details={
                    "task_id": task.task_id,
                    "report_id": report.report_id,
                    "reason": exc.__class__.__name__,
                },
            ) from exc

        self.artifact_repository.save(
            WORD_REPORT_ARTIFACT_TYPE,
            word_report.word_report_id,
            word_report,
        )
        return word_report

    def _cached_word_report(self, task_id: str, report_id: str) -> WordReport | None:
        cached_reports = self.artifact_repository.list_by_task(
            task_id,
            WORD_REPORT_ARTIFACT_TYPE,
            WordReport,
        )
        for cached in reversed(cached_reports):
            word_report = WordReport.model_validate(cached)
            if word_report.report_id != report_id:
                continue
            if word_report.metadata.get("render_version") != WORD_REPORT_RENDER_VERSION:
                continue
            if Path(word_report.file_path).exists():
                return word_report
        return None

    def _latest_cached_report(self, task_id: str) -> ReportData | None:
        cached_reports = self.artifact_repository.list_by_task(
            task_id,
            REPORT_ARTIFACT_TYPE,
            ReportData,
        )
        if not cached_reports:
            return None
        return redact_report_data(ReportData.model_validate(cached_reports[-1]))

    def _get_completed_task(self, task_id: str) -> AnalysisTask:
        task = self.task_repository.get(task_id)
        if task is None:
            raise WordReportServiceError(
                "TASK_NOT_FOUND",
                "Task not found",
                status_code=404,
                details={"task_id": task_id},
            )
        if task.status not in _WORD_REPORT_READABLE_STATUSES:
            raise WordReportServiceError(
                "WORD_REPORT_NOT_READY",
                "Word report is only available after completion or human review.",
                status_code=409,
                details={"task_id": task_id, "status": task.status.value},
            )
        return task

    def _run_workflow(self, task: AnalysisTask) -> Mapping[str, Any]:
        try:
            workflow = self.workflow_factory()
            state = create_initial_state(task)
            result = workflow.invoke(state)
        except Exception as exc:
            raise WordReportServiceError(
                "WORD_REPORT_DATA_FAILED",
                "Word report data generation failed",
                status_code=500,
                details={"task_id": task.task_id, "reason": exc.__class__.__name__},
            ) from exc

        if result["task"].get("status") != TaskStatus.COMPLETED.value or not result["reports"]:
            raise WordReportServiceError(
                "WORD_REPORT_DATA_FAILED",
                "Word report data generation did not produce a completed report.",
                status_code=500,
                details={
                    "task_id": task.task_id,
                    "workflow_status": result["task"].get("status"),
                },
            )
        return result

    def _report_from_state(self, task: AnalysisTask, state: Mapping[str, Any]) -> ReportData:
        report = redact_report_data(ReportData.model_validate(state["reports"][-1]))
        cached = self.artifact_repository.get(
            task.task_id,
            REPORT_ARTIFACT_TYPE,
            report.report_id,
            ReportData,
        )
        if cached is None:
            self.artifact_repository.save(REPORT_ARTIFACT_TYPE, report.report_id, report)
            return report
        return redact_report_data(ReportData.model_validate(cached))

    def _render_relationship_graph(
        self,
        task: AnalysisTask,
        report: ReportData,
        battlefield: BattlefieldData,
    ) -> tuple[Path | None, JsonObject]:
        file_name = f"{task.task_id}_{report.report_id}_competition_graph.png"
        try:
            graph_path = self.graph_renderer(
                battlefield,
                output_dir=self.output_dir,
                file_name=file_name,
            )
        except Exception as exc:
            return None, {
                "relationship_graph_status": "failed",
                "relationship_graph_failure_reason": exc.__class__.__name__,
            }

        graph_image = RelationshipGraphImage(
            graph_image_id=f"relationship_graph_{report.report_id}",
            task_id=task.task_id,
            report_id=report.report_id,
            generated_at=datetime.now(UTC),
            file_path=str(graph_path),
            file_name=graph_path.name,
            byte_size=graph_path.stat().st_size,
            metadata={
                "relation_count": len(battlefield.key_relations[:5]),
                "security_scan": "passed",
                "source": "word_report_export",
            },
        )
        self.artifact_repository.save(
            RELATIONSHIP_GRAPH_ARTIFACT_TYPE,
            graph_image.graph_image_id,
            graph_image,
        )
        return graph_path, {
            "relationship_graph_status": "generated",
            "relationship_graph_image_id": graph_image.graph_image_id,
            "relationship_graph_file_name": graph_image.file_name,
        }

    def _record_word_export_failure(
        self,
        *,
        task_id: str,
        report: ReportData,
        phase: str,
        exc: Exception,
    ) -> None:
        task = self.task_repository.get(task_id)
        if task is None:
            return

        cached = self.artifact_repository.get(
            task_id,
            TRACE_ARTIFACT_TYPE,
            _trace_artifact_id(task_id),
            TraceData,
        )
        if cached is None:
            trace = _build_trace_data(
                task=task,
                state=None,
                trace_view_id=_trace_artifact_id(task_id),
            )
        else:
            trace = TraceData.model_validate(cached)

        failure_record = {
            "status": "failed",
            "code": "WORD_REPORT_EXPORT_FAILED",
            "report_id": report.report_id,
            "phase": phase,
            "error_type": exc.__class__.__name__,
            "readable_reason": "Word 报告导出失败，请检查导出目录或文件写入权限。",
            "details": redact_sensitive_value(
                {
                    "task_id": task_id,
                    "report_id": report.report_id,
                    "phase": phase,
                },
                redact_key_names=True,
            ),
            "recorded_at": datetime.now(UTC).isoformat(),
        }
        metadata = dict(trace.metadata)
        failure_records = list(metadata.get("word_export_failures", []))
        failure_records.append(failure_record)
        metadata["word_export_failures"] = failure_records
        metadata["last_failure"] = failure_record
        updated_trace = TraceData.model_validate(
            {
                **trace.model_dump(mode="json"),
                "metadata": metadata,
                "generated_at": datetime.now(UTC).isoformat(),
            }
        )
        self.artifact_repository.save(
            TRACE_ARTIFACT_TYPE,
            updated_trace.trace_view_id,
            updated_trace,
        )


def render_word_report(
    report_data: ReportData | Mapping[str, Any],
    *,
    products: Sequence[Product | Mapping[str, Any]] = (),
    battlefield: BattlefieldData | Mapping[str, Any] | None = None,
    relationship_graph_path: Path | str | None = None,
    output_dir: Path | str | None = None,
    file_name: str | None = None,
    generated_at: datetime | None = None,
    extra_metadata: Mapping[str, Any] | None = None,
) -> WordReport:
    report = ReportData.model_validate(report_data)
    exported_at = generated_at or datetime.now(UTC)

    document = Document()
    _set_core_properties(document, report)
    _append_cover(document, report, exported_at)
    _append_static_toc(document, report)
    _append_report_body(document, report)
    _assert_document_is_safe(document)

    report_dir = Path(output_dir) if output_dir is not None else DEFAULT_REPORTS_DIR
    report_dir.mkdir(parents=True, exist_ok=True)
    output_path = report_dir / _safe_docx_file_name(
        file_name or f"{report.task_id}_{report.report_id}.docx"
    )
    document.save(output_path)

    metadata: JsonObject = {
        "section_count": len(report.section_order),
        "file_name": output_path.name,
        "byte_size": output_path.stat().st_size,
        "render_version": WORD_REPORT_RENDER_VERSION,
        "security_scan": "passed",
        "target_image_status": "omitted",
        "core_competitor_image_count": 0,
        "relationship_graph_included": False,
    }
    if extra_metadata:
        metadata.update(redact_sensitive_value(dict(extra_metadata), redact_key_names=True))

    return WordReport(
        word_report_id=f"word_{report.report_id}",
        task_id=report.task_id,
        report_id=report.report_id,
        generated_at=exported_at,
        file_path=str(output_path),
        file_name=output_path.name,
        byte_size=output_path.stat().st_size,
        metadata=metadata,
    )


def _set_core_properties(document, report: ReportData) -> None:
    document.core_properties.title = _safe_text("竞品分析报告")
    document.core_properties.subject = _safe_text("自动猫砂盆竞品分析")
    document.core_properties.keywords = "competitive-analysis,docx"


def _append_cover(document, report: ReportData, exported_at: datetime) -> None:
    document.add_heading(_safe_text("自动猫砂盆竞品分析报告"), 0)
    document.add_paragraph(
        _safe_text("本报告面向产品和运营决策阅读，只保留结论、分析理由和行动建议。")
    )
    document.add_paragraph(_safe_text(f"报告生成时间：{_format_datetime(report.generated_at)}"))
    document.add_paragraph(_safe_text(f"Word 导出时间：{_format_datetime(exported_at)}"))
    document.add_page_break()


def _append_static_toc(document, report: ReportData) -> None:
    document.add_heading(_safe_text("目录"), level=1)
    document.add_paragraph(_safe_text("正文"))
    narrative_sections = _narrative_sections(report)
    if narrative_sections:
        for index, section in enumerate(narrative_sections, start=1):
            toc_line = f"{index}. {_narrative_title(section)}"
            document.add_paragraph(_safe_text(toc_line), style="List Number")
        document.add_page_break()
        return
    for index, section in enumerate(_ordered_sections(report), start=1):
        toc_line = f"{index}. {_section_title(section)}"
        document.add_paragraph(_safe_text(toc_line), style="List Number")
    document.add_page_break()


def _append_image_summary(
    document,
    products: Sequence[Product],
    battlefield: BattlefieldData | None,
) -> None:
    document.add_heading(_safe_text("产品图片摘要"), level=1)
    target_product = _target_product(products)
    _append_product_image_block(document, "目标产品缩略图", target_product)

    products_by_id = {product.product_id: product for product in products}
    competitor_ids = []
    if battlefield is not None:
        competitor_ids = [
            relation.competitor_product_id for relation in battlefield.key_relations[:3]
        ]
    document.add_heading(_safe_text("核心竞品缩略图"), level=2)
    if not competitor_ids:
        document.add_paragraph(_safe_text(NO_RELIABLE_IMAGE))
    for competitor_id in competitor_ids:
        competitor = products_by_id.get(competitor_id)
        _append_product_image_block(
            document,
            competitor.name if competitor is not None else competitor_id,
            competitor,
        )


def _append_product_image_block(document, title: str, product: Product | None) -> None:
    document.add_heading(_safe_text(title), level=2)
    if product is None:
        document.add_paragraph(_safe_text(NO_RELIABLE_IMAGE))
        return

    document.add_paragraph(_safe_text(product.name))
    image_path = _image_path_for_product(product)
    if image_path is None:
        document.add_paragraph(_safe_text(NO_RELIABLE_IMAGE))
        return
    if not _add_picture_or_placeholder(document, image_path, width_inches=1.35):
        document.add_paragraph(_safe_text(NO_RELIABLE_IMAGE))


def _append_relationship_graph(document, relationship_graph_path: Path | str | None) -> None:
    document.add_heading(_safe_text("简化竞争关系图"), level=1)
    image_path = _valid_image_path(relationship_graph_path)
    if image_path is None:
        document.add_paragraph(_safe_text(NO_RELIABLE_IMAGE))
        return
    if not _add_picture_or_placeholder(document, image_path, width_inches=6.2):
        document.add_paragraph(_safe_text(NO_RELIABLE_IMAGE))


def _append_report_body(document, report: ReportData) -> None:
    document.add_heading(_safe_text("正文"), level=1)
    narrative_sections = _narrative_sections(report)
    if narrative_sections:
        for section in narrative_sections:
            _append_narrative_section(document, section)
        return
    for section in _ordered_sections(report):
        _append_section(document, section)


def _append_narrative_section(document, section: Mapping[str, Any]) -> None:
    document.add_heading(_safe_text(_narrative_title(section)), level=1)
    paragraphs = _narrative_paragraphs(section)
    if not paragraphs:
        document.add_paragraph(_safe_text(NO_RELIABLE_DATA))
        return
    for paragraph in paragraphs:
        document.add_paragraph(_safe_text(paragraph))


def _append_section(document, section: ReportSection) -> None:
    document.add_heading(_safe_text(_section_title(section)), level=1)
    document.add_paragraph(_safe_text(_section_summary(section)))
    if not section.items:
        document.add_paragraph(_safe_text(NO_RELIABLE_DATA))
        return

    visible_items = section.items[:_MAX_WORD_SECTION_ITEMS]
    for index, item in enumerate(visible_items, start=1):
        _append_report_item(document, section, item, index)
    if len(section.items) > len(visible_items):
        document.add_paragraph(
            _safe_text(
                f"本节另有 {len(section.items) - len(visible_items)} 条结构化关系已纳入分析，"
                "Word 版只展开最需要阅读的判断，完整证据链可在网页端继续查看。"
            )
        )


def _append_report_item(
    document,
    section: ReportSection,
    item: Mapping[str, Any],
    index: int,
) -> None:
    title = _word_item_title(section.section_id, item, index)
    document.add_heading(_safe_text(title), level=2)
    for paragraph in _word_item_paragraphs(section.section_id, item):
        document.add_paragraph(_safe_text(paragraph))


def _word_item_paragraphs(section_id: str, item: Mapping[str, Any]) -> list[str]:
    expanded_paragraphs = _expanded_paragraphs(item)
    if expanded_paragraphs:
        return expanded_paragraphs

    llm_paragraphs = _llm_paragraphs(item)
    if llm_paragraphs:
        return llm_paragraphs

    competitor = _competitor_name(item)
    slice_label = _slice_label(item)
    score = _score_phrase(item.get("edge_score") or item.get("top_edge_score"))
    stages = _decision_stage_text(item.get("decision_stages") or item.get("decision_stage"))
    recommendation = _string_value(item.get("recommendation"))
    if section_id == "core_competitor_analysis" and _string_value(
        item.get("why_users_compare")
    ):
        strengths = _string_list(item.get("competitor_strengths"))
        weaknesses = _string_list(item.get("competitor_weaknesses"))
        return [
            f"为什么是竞品：{_string_value(item.get('why_users_compare'))}",
            f"竞品强项：{_join_cn(strengths) if strengths else NO_RELIABLE_DATA}",
            (
                "我方回应："
                f"{_string_value(item.get('target_response')) or '建议先补齐证据后再判断回应。'}"
            ),
            (
                f"风险边界：{_join_cn(weaknesses)}"
                if weaknesses
                else (
                    "应答话术："
                    f"{_string_value(item.get('response_talk_track')) or '证据不足处建议复核。'}"
                )
            ),
        ]
    if section_id == "target_opportunities_and_risks" and _string_value(
        item.get("dimension")
    ):
        return [
            (
                f"差距维度：{_string_value(item.get('dimension'))}。"
                f"{_string_value(item.get('target_status')) or NO_RELIABLE_DATA}"
            ),
            f"决策影响：{_string_value(item.get('impact_on_decision')) or NO_RELIABLE_DATA}",
            (
                "下一步："
                f"{_string_value(item.get('recommendation')) or '建议补齐证据后再做确定判断。'}"
            ),
        ]
    if section_id == "product_strategy_recommendations" and _string_value(
        item.get("opportunity_id")
    ):
        return [
            (
                f"机会：{_string_value(item.get('title')) or '未命名机会'}。"
                f"{recommendation or '建议围绕核心竞品比较点优化表达。'}"
            ),
            (
                f"预期影响：{_string_value(item.get('expected_impact')) or NO_RELIABLE_DATA}。"
                f"责任方向：{_string_value(item.get('owner')) or NO_RELIABLE_DATA}。"
            ),
            f"证据边界：{_string_value(item.get('evidence_boundary')) or '证据不足处建议复核。'}",
        ]
    if section_id == "competitive_landscape_judgment" and _string_value(
        item.get("competition_meaning")
    ):
        return [
            _string_value(item.get("competition_meaning")),
            _string_value(item.get("why_now")) or "该切片适合优先阅读。",
        ]
    if section_id in {"conclusion_summary", "core_competitor_analysis"}:
        return [
            (
                f"{competitor}是当前最需要优先解释的竞品对象。它与目标产品争夺的是"
                "相近的购买理由：用户希望减少清理负担、控制异味，并确认长期维护成本是否可接受。"
            ),
            (
                f"从现有证据看，这条关系{score}。"
                f"{f'放在{slice_label}下看，' if slice_label else ''}"
                "用户不会只比较单一参数，而是会同时比较清理省心程度、容量、除臭可信度和价格解释。"
            ),
            (
                "建议下一步把目标产品的差异讲成用户能直接理解的语言：为什么更省心、"
                "为什么更可信、为什么价格值得，而不是堆功能名或内部评分。"
            ),
        ]
    if section_id in {"competitive_landscape_judgment", "dynamic_slice_analysis"}:
        return [
            (
                f"{slice_label or '当前重点切片'}下的竞争压力主要来自{competitor}。"
                "这个切片说明用户会把目标产品和相近方案放在同一组候选里比较。"
            ),
            (
                "竞争焦点不是“有没有自动清理”这样单一功能，而是谁能更完整地解释"
                "自动清理、除臭、容量和维护成本之间的取舍。"
            ),
            "建议报告阅读时优先看该切片下的购买阻力：如果用户无法快速理解差异，就会自然转向表达更清楚或价格更容易接受的竞品。",
        ]
    if section_id in {"user_decision_chain_analysis", "decision_chain_analysis"}:
        return [
            f"在{stages}阶段，用户最关心的是产品能力是否可靠、使用风险是否可控，以及售后和维护是否省心。",
            (
                "当前竞争关系会影响用户是否继续把目标产品留在候选集中。若目标产品只强调功能名，"
                "却没有解释真实使用收益，用户很容易转向更好理解的竞品。"
            ),
            "建议把这一阶段的表达从参数说明改成购买决策语言，例如清理频率能减少多少、异味控制如何被验证、维护成本为什么可接受。",
        ]
    if section_id == "target_opportunities_and_risks" and _string_value(item.get("dimension")):
        return _string_value(item.get("dimension"))
    if section_id == "product_strategy_recommendations" and _string_value(item.get("title")):
        return _string_value(item.get("title"))
    if section_id in {"product_strategy_recommendations", "recommendations"}:
        return [
            recommendation
            or (
                "当前最重要的策略动作，是把目标产品与核心竞品之间的差异讲清楚，"
                "并补足价格、除臭、安全或维护成本相关证据。"
            ),
            "建议优先处理会直接影响购买决策的信息：核心竞品对比、用户异议回应、证据可信度说明和下一步复核事项。",
            "不要把内部评分或字段解释交给用户阅读，应把它们转化成清楚的结论、原因和行动建议。",
        ]
    if section_id in {"target_opportunities_and_risks", "product_profile"}:
        return _profile_paragraphs(item)
    if section_id in {"evidence_quality_appendix", "analysis_process_appendix", "evidence_index"}:
        return [
            "本节只保留证据与流程的业务含义，不展开内部字段、任务编号或原始截图路径。",
            "证据不足的内容应保持保守表达；如果后续补充了更完整的评论、销量、价格或认证材料，相关结论可以继续更新。",
        ]
    return _readable_mapping_paragraphs(item)


def _section_summary(section: ReportSection) -> str:
    match section.section_id:
        case "conclusion_summary":
            return (
                "本节先给出最重要的总体判断：目标产品面对谁的压力，"
                "以及用户为什么会把它们放在一起比较。"
            )
        case "competitive_landscape_judgment" | "dynamic_slice_analysis":
            return "本节按价格带、人群和使用场景梳理竞争压力，只展开最值得优先看的切片。"
        case "core_competitor_analysis" | "competitor_findings":
            return "本节拆解最容易被用户拿来横向比较的核心竞品，重点看购买理由和转化阻力。"
        case "user_decision_chain_analysis" | "decision_chain_analysis":
            return "本节说明用户在不同购买阶段会如何改变选择，以及目标产品需要补强的表达。"
        case "target_opportunities_and_risks" | "product_profile":
            return "本节总结目标产品已经具备的表达机会，以及当前仍需要保守处理或补证的风险。"
        case "product_strategy_recommendations" | "recommendations":
            return "本节把前文判断转化成可执行的内容表达、证据补充和产品策略建议。"
        case "evidence_quality_appendix":
            return "本节保留证据质量的业务结论，不展示内部审计字段。"
        case "analysis_process_appendix":
            return "本节只简要说明分析流程，详细 Trace 和技术日志请在网页端查看。"
        case _:
            return section.summary


def _word_item_title(section_id: str, item: Mapping[str, Any], index: int) -> str:
    if section_id in {"competitive_landscape_judgment", "dynamic_slice_analysis"}:
        return f"重点切片 {index}：{_slice_label(item) or '当前竞争场景'}"
    if section_id in {"core_competitor_analysis", "conclusion_summary"}:
        return f"核心竞品 {index}：{_competitor_name(item)}"
    if section_id in {"user_decision_chain_analysis", "decision_chain_analysis"}:
        return f"决策阶段 {index}：{_decision_stage_text(item.get('decision_stage'))}"
    if section_id in {"product_strategy_recommendations", "recommendations"}:
        return f"行动建议 {index}"
    if section_id in {"target_opportunities_and_risks", "product_profile"}:
        return "目标产品机会与风险"
    return f"分析要点 {index}"


def _llm_paragraphs(item: Mapping[str, Any]) -> list[str]:
    paragraphs = item.get("llm_paragraphs")
    if not isinstance(paragraphs, Mapping):
        return []
    values = []
    for key in ("conclusion", "reason", "action"):
        value = _string_value(paragraphs.get(key))
        if value:
            values.append(value)
    return values


def _expanded_paragraphs(item: Mapping[str, Any]) -> list[str]:
    paragraphs = item.get("llm_expanded_analysis")
    if not isinstance(paragraphs, Sequence) or isinstance(paragraphs, str):
        return []
    return [value for value in (_string_value(item) for item in paragraphs) if value]


def _profile_paragraphs(item: Mapping[str, Any]) -> list[str]:
    product = item.get("product")
    product_name = ""
    if isinstance(product, Mapping):
        product_name = _string_value(product.get("name"))
    insights = item.get("llm_extracted_insights")
    insight_text = _insight_summary(insights)
    return [
        (
            f"{product_name or '目标产品'}的机会在于把“省心清理、除臭可信、"
            "维护成本可控”讲成连续的购买理由，而不是只罗列功能点。"
        ),
        insight_text
        or (
            "当前评论和卖点洞察仍然有限，报告会避免把缺少证据的价格、"
            "销量、安全或认证信息写成确定结论。"
        ),
        "建议后续优先补充评论痛点、真实使用阻力、购买异议和竞品对比证据，让报告从“结构完整”进一步变成“判断有料”。",
    ]


def _insight_summary(value: Any) -> str:
    if not isinstance(value, Sequence) or isinstance(value, str):
        return ""
    parts: list[str] = []
    for insight in value[:3]:
        if not isinstance(insight, Mapping):
            continue
        pain_points = _string_list(insight.get("pain_points"))
        buying_reasons = _string_list(insight.get("buying_reasons"))
        objections = _string_list(insight.get("objections"))
        if pain_points:
            parts.append(f"用户痛点集中在{_join_cn(pain_points)}")
        if buying_reasons:
            parts.append(f"购买理由包括{_join_cn(buying_reasons)}")
        if objections:
            parts.append(f"主要异议是{_join_cn(objections)}")
    return "；".join(parts) + "。" if parts else ""


def _readable_mapping_paragraphs(item: Mapping[str, Any]) -> list[str]:
    readable_parts: list[str] = []
    for key, value in item.items():
        if key in _INTERNAL_REPORT_KEYS or key.endswith("_id") or key.endswith("_ids"):
            continue
        if isinstance(value, Mapping | list | tuple):
            continue
        text = _string_value(value)
        if text:
            readable_parts.append(text)
    if not readable_parts:
        return ["本节暂无更多可直接写入正文的可靠分析内容。"]
    return ["；".join(readable_parts[:4]) + "。"]


def _competitor_name(item: Mapping[str, Any]) -> str:
    competitor = item.get("competitor")
    if isinstance(competitor, Mapping):
        name = _string_value(competitor.get("name"))
        if name:
            return name
    product = item.get("product")
    if isinstance(product, Mapping):
        name = _string_value(product.get("name"))
        if name:
            return name
    return "核心竞品"


def _slice_label(item: Mapping[str, Any]) -> str:
    slice_value = item.get("slice")
    if isinstance(slice_value, Mapping):
        price_band = _string_value(slice_value.get("price_band"))
        persona = _string_value(slice_value.get("persona"))
        scenario = _string_value(slice_value.get("scenario"))
    else:
        price_band = _string_value(item.get("price_band"))
        persona = _string_value(item.get("persona"))
        scenario = _string_value(item.get("scenario"))
    return " / ".join(value for value in (price_band, persona, scenario) if value)


def _score_phrase(value: Any) -> str:
    if isinstance(value, int | float):
        return f"竞争强度约为 {value:.0%}"
    return "属于需要重点关注的竞争关系"


def _decision_stage_text(value: Any) -> str:
    if isinstance(value, Sequence) and not isinstance(value, str):
        stages = [_string_value(item) for item in value if _string_value(item)]
        return _join_cn(stages) if stages else "购买决策"
    return _string_value(value) or "购买决策"


def _string_value(value: Any) -> str:
    if value is None or value == "" or value == [] or value == {}:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, float):
        return f"{value:.2f}"
    text = str(value).strip()
    if _looks_internal_value(text):
        return ""
    return text


def _string_list(value: Any) -> list[str]:
    if not isinstance(value, Sequence) or isinstance(value, str):
        return []
    return [item for item in (_string_value(item) for item in value) if item][:4]


def _join_cn(values: Sequence[str]) -> str:
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    return "、".join(values[:-1]) + f"和{values[-1]}"


def _looks_internal_value(value: str) -> bool:
    return bool(
        re.match(r"^(task|run|edge|claim|ev|prod|review|trace|tool|msg)_[A-Za-z0-9_]+$", value)
        or re.search(r"\b(edge|claim|evidence|task|product)_id\b", value, flags=re.IGNORECASE)
    )


def _format_datetime(value: datetime) -> str:
    return value.strftime("%Y/%m/%d %H:%M")


def _add_picture_or_placeholder(document, image_path: Path, *, width_inches: float) -> bool:
    try:
        document.add_picture(str(image_path), width=Inches(width_inches))
    except Exception:
        return False
    return True


def _target_product(products: Sequence[Product]) -> Product | None:
    for product in products:
        if product.role == ProductRole.TARGET:
            return product
    return products[0] if products else None


def _core_competitor_image_count(
    products: Sequence[Product],
    battlefield: BattlefieldData | None,
) -> int:
    if battlefield is None:
        return 0
    products_by_id = {product.product_id: product for product in products}
    count = 0
    for relation in battlefield.key_relations[:3]:
        product = products_by_id.get(relation.competitor_product_id)
        if product is not None and _image_path_for_product(product) is not None:
            count += 1
    return count


def _image_status(product: Product | None) -> str:
    if product is None:
        return "missing"
    return "available" if _image_path_for_product(product) is not None else "missing"


def _image_path_for_product(product: Product) -> Path | None:
    for value in (
        product.primary_image_source_path,
        product.primary_image_path,
        product.primary_image_url,
    ):
        candidate = _resolve_local_image_path(value)
        if candidate is not None:
            return candidate
    return None


def _resolve_local_image_path(value: str | None) -> Path | None:
    if value is None or not value.strip():
        return None
    raw_value = value.strip()
    if raw_value.startswith(("http://", "https://")):
        return None
    if raw_value.startswith(_RAW_ASSET_URL_PREFIX):
        relative = unquote(raw_value.removeprefix(_RAW_ASSET_URL_PREFIX))
        return _valid_image_path(_RAW_ASSETS_DIR / relative, root=_RAW_ASSETS_DIR)
    path = Path(raw_value)
    if not path.is_absolute():
        path = _PROJECT_ROOT / path
    return _valid_image_path(path, root=_PROJECT_ROOT)


def _valid_image_path(value: Path | str | None, *, root: Path | None = None) -> Path | None:
    if value is None:
        return None
    path = Path(value)
    if path.suffix.lower() not in _SUPPORTED_DOCX_IMAGE_EXTENSIONS:
        return None
    try:
        resolved = path.resolve()
    except OSError:
        return None
    if root is not None:
        try:
            resolved.relative_to(root.resolve())
        except ValueError:
            return None
    if not resolved.exists() or not resolved.is_file():
        return None
    return resolved


def _narrative_sections(report: ReportData) -> list[JsonObject]:
    narrative_report = report.narrative_report
    if not isinstance(narrative_report, Mapping):
        return []
    sections = narrative_report.get("sections")
    if not isinstance(sections, Sequence) or isinstance(sections, str):
        return []
    result: list[JsonObject] = []
    for section in sections:
        if not isinstance(section, Mapping):
            continue
        paragraphs = _narrative_paragraphs(section)
        if not paragraphs:
            continue
        result.append(
            {
                "section_id": _string_value(section.get("section_id")),
                "title": _narrative_title(section),
                "paragraphs": paragraphs,
            }
        )
    return result


def _narrative_title(section: Mapping[str, Any]) -> str:
    title = _string_value(section.get("title"))
    if title:
        return title
    section_id = _string_value(section.get("section_id"))
    return {
        "executive_summary": "执行摘要",
        "competitive_landscape": "竞争格局",
        "core_competitors": "核心竞品",
        "decision_chain": "用户决策链",
        "action_recommendations": "行动建议",
    }.get(section_id, "分析章节")


def _narrative_paragraphs(section: Mapping[str, Any]) -> list[str]:
    paragraphs = section.get("paragraphs")
    if not isinstance(paragraphs, Sequence) or isinstance(paragraphs, str):
        return []
    readable: list[str] = []
    for paragraph in paragraphs:
        text = _string_value(paragraph)
        if not text or _looks_internal_value(text):
            continue
        readable.append(text)
    return readable[:6]


def _ordered_sections(report: ReportData) -> list[ReportSection]:
    return [getattr(report, section_id) for section_id in report.section_order]


def _section_title(section: ReportSection) -> str:
    modern_titles = {
        "conclusion_summary": "执行摘要",
        "competitive_landscape_judgment": "竞争格局",
        "core_competitor_analysis": "核心竞品 Battlecard",
        "user_decision_chain_analysis": "用户决策链",
        "target_opportunities_and_risks": "差距矩阵",
        "product_strategy_recommendations": "机会地图与优先级",
        "evidence_quality_appendix": "风险与证据边界",
        "analysis_process_appendix": "附录",
    }
    return modern_titles.get(
        section.section_id,
        _SECTION_TITLE_OVERRIDES.get(section.section_id, section.title),
    )


def _assert_document_is_safe(document) -> None:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if contains_sensitive_text(text):
        raise WordRenderError("Word report export blocked by sensitive content scan.")


def _safe_docx_file_name(value: str) -> str:
    safe_value = redact_sensitive_text(Path(value).name)
    stem = _SAFE_FILE_STEM_PATTERN.sub("_", Path(safe_value).stem).strip("._")
    return f"{stem or 'report'}.docx"


def _safe_text(value: Any) -> str:
    redacted = redact_sensitive_text(str(value or ""))
    compact = " ".join(redacted.split())
    if len(compact) <= _MAX_TEXT_CHARS:
        return compact
    return compact[: _MAX_TEXT_CHARS - 3].rstrip() + "..."


def _format_scalar(value: Any) -> str:
    if value is None or value == "" or value == [] or value == {}:
        return NO_RELIABLE_DATA
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def _humanize_key(key: str) -> str:
    return key.replace("_", " ").title()


def _list_style(depth: int) -> str:
    return "List Bullet 2" if depth > 0 else "List Bullet"
